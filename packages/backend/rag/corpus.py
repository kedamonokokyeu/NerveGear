"""
rag/corpus.py

Source documents for the demo index.

load_docx() pulls plain text from a .docx so the assistant can answer questions
about documents like the NerveGear roadmap. cfd_glossary() is a small built-in set of
fluid-dynamics definitions so the pipeline has useful content even with no files.
Both return Document objects ready for RagPipeline.index_documents().
"""

from __future__ import annotations

from .pipeline import Document


def load_docx(path: str, doc_id: str = "roadmap") -> Document:
    """Extract text from a .docx. Requires python-docx."""
    from docx import Document as Docx  # lazy

    d = Docx(path)
    text = "\n\n".join(p.text for p in d.paragraphs if p.text.strip())
    return Document(id=doc_id, text=text, source=path)


def cfd_glossary() -> list[Document]:
    """
    Mini knowledge base so demo can work without external files.
    Later on, make default so that RAG pipeline always retrieves something even if user doesn't upload files.
    """
    entries = {
        "reynolds_number": (
            "The Reynolds number (Re) is a dimensionless quantity that measures the "
            "ratio of inertial forces to viscous forces in a fluid flow, Re = rho*U*L/mu. "
            "Low Re (<~2000 in a pipe) is laminar and smooth; high Re is turbulent and "
            "chaotic. In NerveGear's 2D surrogate the training range is roughly 50 to 400, "
            "so flows stay in the laminar-to-transitional regime."
        ),
        "inlet_velocity": (
            "Inlet velocity is the speed and direction of fluid entering the domain "
            "boundary. In NerveGear it is set by the inlet_edge (left/right/top/bottom), a "
            "speed magnitude (0.05 to 0.15 in trained units), and an angle (+/- 30 "
            "degrees) tilting the flow off the edge normal."
        ),
        "pressure_field": (
            "The pressure field is the scalar pressure at every point in the domain. "
            "Gradients in pressure drive the flow; high pressure builds on the upstream "
            "face of an obstacle (stagnation) and drops in accelerated regions and wakes."
        ),
        "signed_distance_field": (
            "A signed distance field (SDF) stores, at each grid point, the distance to "
            "the nearest surface, negative inside the solid and positive outside. NerveGear "
            "feeds an SDF (plus a solid mask) to the network so it knows the geometry; "
            "the SDF gives smooth gradients that help the model localize boundaries."
        ),
        "navier_stokes": (
            "The Navier-Stokes equations are the partial differential equations governing "
            "viscous fluid motion, expressing conservation of momentum and mass. Solving "
            "them directly (as ANSYS/OpenFOAM do) is expensive; NerveGear trains a neural "
            "surrogate to predict the resulting velocity and pressure fields directly, and "
            "validates against these equations."
        ),
        "surrogate_model": (
            "A surrogate model is a fast approximation that learns the input-to-output "
            "mapping of an expensive simulator. NerveGear's surrogate is a U-Net taking a "
            "5-channel input (SDF, mask, Reynolds map, inlet vx, inlet vy) and predicting "
            "3 channels (vx, vy, pressure) on a 128x128 grid in milliseconds."
        ),
    }
    return [Document(id=k, text=v, source=f"cfd_glossary/{k}") for k, v in entries.items()]


def nervegear_corpus(repo_root: str | None = None) -> list[Document]:
    """
    Everything the assistant should know: the built-in glossary, the
    microfluidics knowledge base, the repo docs (contract, pivots, README),
    and the roadmap .docx when present. Missing files are skipped silently so
    the pipeline works in any checkout state.
    """
    import os

    here = os.path.dirname(os.path.abspath(__file__))
    root = repo_root or os.path.normpath(os.path.join(here, "..", "..", "..", ".."))
    docs = list(cfd_glossary())

    def add_text(path, doc_id):
        try:
            with open(path, encoding="utf-8") as f:
                docs.append(Document(id=doc_id, text=f.read(), source=os.path.basename(path)))
        except OSError:
            pass

    add_text(os.path.join(here, "knowledge", "microfluidics.md"), "microfluidics")
    add_text(os.path.join(root, "docs", "CONTRACT.md"), "contract")
    add_text(os.path.join(root, "README.md"), "readme")
    add_text(os.path.join(root, "PROJECT_MAP.md"), "project_map")
    ax3d = os.path.normpath(os.path.join(here, ".."))
    add_text(os.path.join(ax3d, "MICROSCALE_PIVOT.md"), "microscale_pivot")
    add_text(os.path.join(ax3d, "THERMAL.md"), "thermal")
    add_text(os.path.join(ax3d, "MODEL_CARD.md"), "model_card")
    roadmap = os.path.join(root, "docs", "NerveGear_Roadmap.docx")
    if os.path.exists(roadmap):
        try:
            docs.append(load_docx(roadmap, "roadmap"))
        except Exception:  # noqa: BLE001 - python-docx missing or file corrupt
            pass
    return docs
